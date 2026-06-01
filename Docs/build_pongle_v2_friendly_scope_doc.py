from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFilter, ImageFont
from pathlib import Path


OUT = "Docs/Pongle_V2_Friendly_Scope.docx"
LOGO = "Pongle/Assets.xcassets/BatOrange.imageset/BatOrange.png"
DIAGRAM = "Docs/Pongle_V2_Architecture.png"

ORANGE = RGBColor(255, 92, 28)
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(255, 255, 255)
INK = RGBColor(24, 24, 24)
MUTED = RGBColor(92, 92, 92)
BLACK_FILL = "111111"
SOFT = "FFFFFF"


def load_font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/SFNS.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def wrap_text(draw, text, font, max_width):
    lines = []
    for raw_line in text.split("\n"):
        words = raw_line.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            trial = f"{current} {word}"
            width = draw.textbbox((0, 0), trial, font=font)[2]
            if width <= max_width:
                current = trial
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def draw_centered_text(draw, box, text, font, fill, line_gap=6):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, x2 - x1)
    heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total_h = sum(heights) + line_gap * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line, height in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += height + line_gap


def draw_card(draw, shadow, box, title, body, fill, outline, title_fill="#111111", body_fill="#343434"):
    x1, y1, x2, y2 = box
    radius = 34
    draw.rounded_rectangle((x1 + 10, y1 + 14, x2 + 10, y2 + 14), radius=radius, fill="#EDEDED")
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=4)
    title_font = load_font(38, bold=True)
    body_font = load_font(27)
    draw_centered_text(draw, (x1 + 28, y1 + 28, x2 - 28, y1 + 82), title, title_font, title_fill, line_gap=4)
    draw_centered_text(draw, (x1 + 34, y1 + 92, x2 - 34, y2 - 24), body, body_font, body_fill, line_gap=8)


def draw_arrow(draw, start, end, color="#FF5C1C", width=7):
    import math

    draw.line((start, end), fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    head_len = 28
    head_angle = math.pi / 7
    points = [
        end,
        (
            end[0] - head_len * math.cos(angle - head_angle),
            end[1] - head_len * math.sin(angle - head_angle),
        ),
        (
            end[0] - head_len * math.cos(angle + head_angle),
            end[1] - head_len * math.sin(angle + head_angle),
        ),
    ]
    draw.polygon(points, fill=color)


def generate_architecture_diagram():
    Path(DIAGRAM).parent.mkdir(parents=True, exist_ok=True)
    w, h = 1800, 1040
    image = Image.new("RGB", (w, h), "white")
    shadow = Image.new("RGBA", (w, h), (0, 0, 0, 0))
    draw = ImageDraw.Draw(image)
    shadow_draw = ImageDraw.Draw(shadow)

    orange = "#FF5C1C"
    black = "#111111"
    ink = "#242424"
    muted = "#6B6B6B"
    light_orange = "#FFF4EE"
    light_gray = "#F6F6F6"

    title_font = load_font(54, bold=True)
    sub_font = load_font(30)
    draw.text((90, 70), "How Pongle V2 fits together", font=title_font, fill=black)
    draw.text(
        (92, 138),
        "A simple online layer around the scorekeeper: players, places, results, and confirmation.",
        font=sub_font,
        fill=muted,
    )

    cards = {
        "players": (110, 300, 440, 500),
        "app": (590, 280, 980, 520),
        "backend": (1190, 280, 1620, 520),
        "profiles": (360, 660, 690, 840),
        "tables": (735, 660, 1065, 840),
        "matches": (1110, 660, 1440, 840),
        "sms": (1500, 650, 1710, 850),
    }

    draw_card(draw, shadow_draw, cards["players"], "Players", "play, score, and confirm", light_orange, orange)
    draw_card(draw, shadow_draw, cards["app"], "Pongle App", "Swift app on iPhone, iPad, and Watch", black, black, "white", "#F3F3F3")
    draw_card(draw, shadow_draw, cards["backend"], "Supabase", "secure accounts, database, storage, and sync", light_gray, black)
    draw_card(draw, shadow_draw, cards["profiles"], "Profiles", "phone ID, photo, history, rating", "white", orange)
    draw_card(draw, shadow_draw, cards["tables"], "Tables", "locations, notes, owner, visibility", "white", orange)
    draw_card(draw, shadow_draw, cards["matches"], "Matches", "players, table, date, scores, winner", "white", orange)
    draw_card(draw, shadow_draw, cards["sms"], "SMS", "invite, claim, confirm", light_orange, orange)

    draw_arrow(draw, (440, 400), (590, 400), orange)
    draw_arrow(draw, (980, 400), (1190, 400), orange)
    draw_arrow(draw, (1385, 520), (1385, 600), orange)
    draw.line((535, 600, 1385, 600), fill=orange, width=7)
    draw_arrow(draw, (535, 600), (535, 655), orange)
    draw_arrow(draw, (900, 600), (900, 655), orange)
    draw_arrow(draw, (1385, 600), (1385, 655), orange)
    draw_arrow(draw, (1500, 750), (1448, 750), orange)

    label_font = load_font(24, bold=True)
    draw.text((485, 360), "uses", font=label_font, fill=muted)
    draw.text((1032, 360), "saves + syncs", font=label_font, fill=muted)
    draw.text((750, 560), "what V2 remembers", font=label_font, fill=muted)

    later_box = (360, 900, 1440, 980)
    draw.rounded_rectangle(later_box, radius=24, fill="#FAFAFA", outline="#D8D8D8", width=3)
    later_font = load_font(27, bold=True)
    draw_centered_text(
        draw,
        (later_box[0] + 28, later_box[1], later_box[2] - 28, later_box[3]),
        "Later, separate scope: clubs, brackets, payments, TV mode, other sports",
        later_font,
        "#555555",
        line_gap=5,
    )

    image.save(DIAGRAM, quality=95)


def set_run_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, after=6, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=14.5, color=BLACK, bold=True)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "FF5C1C")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        shade_cell(cell, BLACK_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.5, bold=True, color=WHITE)

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            if row_idx % 2 == 0:
                shade_cell(cell, "FAFAFA")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(value)
            set_run_font(r, size=9.4)
            if col_idx in (0, 3):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(table, widths)
    add_para(doc, "", after=2)


def add_note_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.left_indent = Pt(12)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=BLACK, bold=True)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "FF5C1C")
    p_bdr.append(left)
    p_pr.append(p_bdr)


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(14)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "FF5C1C")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_at_a_glance(doc):
    rows = [
        (
            "Core V2",
            "Profiles, table locations, match history, opponent selection, and result confirmation.",
        ),
        (
            "Timeline",
            "Target under one month for the focused V2 scope, assuming quick feedback and no wishlist expansion.",
        ),
        (
            "Goal",
            "Help a player find a table, play a match, save the result, and find better future games.",
        ),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    for idx, (label, body) in enumerate(rows):
        cell = table.cell(0, idx)
        shade_cell(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(label)
        set_run_font(r, size=10, color=ORANGE, bold=True)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.05
        r2 = p2.add_run(body)
        set_run_font(r2, size=9.4, color=INK)
    add_para(doc, "", after=2)


def build_doc():
    generate_architecture_diagram()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.text = "Pongle V2"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = "Simple V2 scope draft"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED

    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_p.paragraph_format.space_after = Pt(0)
    logo_p.add_run().add_picture(LOGO, width=Inches(0.52))

    title = add_para(doc, "Pongle V2", size=28, bold=True, after=0)
    title.runs[0].font.color.rgb = BLACK
    add_para(
        doc,
        "Simple scope draft",
        size=14,
        color=ORANGE,
        bold=True,
        after=4,
    )
    add_para(
        doc,
        "A focused next step from the working scorekeeping prototype into profiles, table locations, match history, and early ratings.",
        size=11.5,
        color=MUTED,
        after=4,
    )
    add_divider(doc)
    add_para(doc, "Hi Cody,", after=5)
    add_para(
        doc,
        "Here is how I would think about V2 in a simple and realistic way. V1 proved the scoring experience can work during real games. V2 should add just enough structure so Pongle can remember who played, where they played, what the result was, and start helping people find better games.",
    )
    add_note_box(
        doc,
        "Main idea: keep the scoring experience, then add the simplest online layer needed for players, tables, and match records.",
    )
    add_at_a_glance(doc)

    add_para(
        doc,
        "I am treating the recent scorekeeper polish as V1 work. This draft is focused on V2: the online layer for players, tables, match history, and ratings.",
        color=MUTED,
        italic=True,
        after=10,
    )

    add_heading(doc, "What I Think V2 Should Include")
    rows = [
        ("1", "Player profiles", "Basic player info, phone number as the unique ID, optional photo, match history, and early rating.", "Medium"),
        ("2", "Table locations", "Add real tables with address/map point, notes, visibility, and who added the table.", "Medium"),
        ("3", "Opponent selection", "Choose an existing player or quickly enter a new person before saving a match.", "Medium"),
        ("4", "Match history", "Save match/game results from the scoring flow: players, table, date, game scores, and winner.", "Medium"),
        ("5", "Guest / unclaimed profiles", "If Canon is not on Pongle yet, create his profile from name + phone and let him claim it later.", "Harder"),
        ("6", "SMS invite + consent", "Send the non-user an invite, then let him consent and claim the profile before it becomes a full account.", "Harder"),
        ("7", "Result confirmation", "Let both players confirm the result so ratings are more trustworthy.", "Harder"),
        ("8", "Pongle rating", "Start with a simple rating that feels more fun and approachable than a bland ITTF-style number.", "Harder"),
        ("9", "Share card", "Create a clean result card that can be shared after a match.", "Medium"),
    ]
    add_table(doc, ["Priority", "Feature", "Plain meaning", "Effort"], rows, [780, 1760, 5580, 1240])

    add_heading(doc, "Specific Things I Heard From You")
    decision_rows = [
        ("Phone as ID", "Use phone number as the unique user identifier, with email optional instead of required."),
        ("Unclaimed profile", "Support guest-style profiles first, then let that person claim the profile later."),
        ("Canon use case", "A player can add Canon by first name, last name, and phone number before Canon joins Pongle."),
        ("SMS invite then consent", "Canon receives an invite, reviews/consents, and then claims the profile."),
        ("Duplicate merge", "If duplicate profiles happen, provide a merge path that keeps matches, games, and rating history."),
        ("Ratings tone", "Make Pongle rating feel simple, useful, and more fun than a bland ITTF-style system."),
    ]
    add_table(doc, ["Decision", "How I would handle it"], decision_rows, [2100, 7260])

    add_heading(doc, "Simple V2 Architecture")
    add_para(
        doc,
        "At a high level, V2 adds one online layer around the existing scorekeeper. The app stays simple for players, while Supabase quietly stores the players, tables, matches, and confirmations behind the scenes.",
    )
    diagram = doc.add_paragraph()
    diagram.alignment = WD_ALIGN_PARAGRAPH.CENTER
    diagram.paragraph_format.space_before = Pt(4)
    diagram.paragraph_format.space_after = Pt(6)
    diagram.add_run().add_picture(DIAGRAM, width=Inches(6.45))
    add_para(
        doc,
        "The important idea: users should feel like they are just playing and saving results. The technical pieces should stay mostly invisible.",
        size=10.2,
        color=MUTED,
        italic=True,
        after=8,
    )

    add_heading(doc, "What I Would Save For Later")
    add_para(
        doc,
        "I would not put everything into V2. Tournament brackets, full club management, payments, credits, subscriptions, AirPods input, TV display mode, and pickleball or other sports are all possible later. They are good ideas, but they would make V2 much bigger and more expensive if we include them now.",
    )

    add_heading(doc, "Timeline And Scope Boundary")
    add_para(
        doc,
        "For the full V2 scope below, I would treat this as one fixed-price project once the scope is confirmed. Target timeline: under one month, roughly 3-4 weeks, assuming we keep V2 focused and testing feedback is quick.",
        bold=True,
    )
    add_para(
        doc,
        "That timeline assumes we keep V2 focused: profiles, table locations, opponent selection, saved match history, guest/unclaimed profiles, SMS invite/consent, result confirmation, duplicate merging, a simple Pongle rating, and a share card. It does not include brackets, payments, full club admin, TV app, AirPods input, or other sports.",
    )

    add_heading(doc, "Suggested Build Order")
    phase_rows = [
        ("Step 1", "Profiles, table locations, opponent selection, saved match/game history, and simple profile/table screens.", "1-2 weeks"),
        ("Step 2", "Guest/unclaimed profiles, SMS invite/claim/consent flow, result confirmation, duplicate-profile merge, early Pongle rating, and share cards.", "2-3 weeks"),
        ("Later", "Clubs, brackets, payments, subscriptions, AirPods input, TV display mode, and other sports.", "Separate scope"),
    ]
    add_table(doc, ["Order", "Scope", "Timing"], phase_rows, [1050, 6510, 1800])

    add_heading(doc, "Data Model In Salesforce Terms")
    data_rows = [
        ("Profiles", "first name, last name, phone as unique ID, email, photo, claimed/guest/unclaimed, rating", "A profile plays matches, confirms results, and may add tables."),
        ("Tables", "name, venue, address, map location, visibility, notes, photo, added by", "A table can have many matches and can belong to a venue/club later."),
        ("Matches", "players, table, date/time, match format, winner, confirmation status", "A match belongs to one table and has multiple games."),
        ("Games", "game number, player scores, first server, side order, winner", "Games belong to a match and create the match result."),
        ("Confirmations", "match, player, confirmed/disputed/pending, timestamp", "Confirmations make match history and ratings more trustworthy."),
        ("Merge records", "old profile, new profile, reason, timestamp", "Needed if one person ends up with duplicate profiles."),
    ]
    add_table(doc, ["Object", "Main fields", "Relationship"], data_rows, [1550, 4100, 3710])

    add_heading(doc, "Backend Note")
    add_para(
        doc,
        "I would not build this on top of Salesforce. The Salesforce way of thinking is useful for modeling the data, but Pongle should use a mobile-first backend. My recommendation is Supabase for accounts, database, storage, and sync. Firebase would also be possible, but I would keep the current native Swift iPhone/iPad/Apple Watch app as the front end either way.",
    )

    add_heading(doc, "Ongoing Running Costs")
    add_para(
        doc,
        "These are third-party costs you would pay directly to run Pongle after launch. They are separate from the fixed build cost. Current public prices can change, but this is the realistic starting picture.",
    )
    cost_rows = [
        ("Apple Developer Program", "$99 / year", "Needed to publish, use TestFlight, and distribute the app under your own Apple account."),
        ("Supabase", "Free for early testing; likely $25 / month for production", "Stores profiles, tables, match history, photos, and sync. Usage overages only matter if the app grows."),
        ("SMS invites / phone verification", "Usage-based; usually small at early volume", "Used for invite, claim, and confirmation messages. Cost depends on provider, country, and message volume."),
        ("Domain / email", "Optional, usually low", "Only needed if Pongle wants a website, branded support email, or custom auth/email sender."),
        ("Apple hosting, TestFlight, push notifications, Apple Maps", "Included / no separate basic fee", "Normal App Store/TestFlight distribution and standard Apple platform features do not add a separate server bill."),
    ]
    add_table(doc, ["Cost item", "Expected cost", "Plain meaning"], cost_rows, [2250, 2650, 4460])
    add_note_box(
        doc,
        "Practical estimate: once V2 is live, you should expect roughly $25-$50/month plus the $99/year Apple Developer account, before any large-scale SMS usage or marketing spend.",
    )

    add_heading(doc, "Success For V2")
    add_para(
        doc,
        "V2 is successful if a user can add a real table, choose or enter an opponent, play a match, save the result, see match history, and start building a lightweight Pongle rating profile. That is enough to test with real users without overbuilding the whole long-term platform too early.",
    )

    add_para(
        doc,
        "My recommendation is to keep V2 focused, learn from real players, and then decide what the next layer should be.",
        bold=True,
        color=ORANGE,
        after=0,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
