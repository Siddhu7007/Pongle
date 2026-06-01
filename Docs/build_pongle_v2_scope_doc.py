from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Docs/Pongle_V2_Scope.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(34, 34, 34)
MUTED = RGBColor(89, 89, 89)
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = width_to_inches(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def width_to_inches(dxa):
    return Inches(dxa / 1440)


def set_run_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        size, color, before, after = 16, BLUE, 16, 8
    elif level == 2:
        size, color, before, after = 13, BLUE, 12, 6
    else:
        size, color, before, after = 12, DARK_BLUE, 8, 4
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=True)
    return p


def add_title_block(doc):
    add_para(doc, "Pongle V2 Scope Brief", size=23, bold=True, after=4)
    add_para(
        doc,
        "Purpose: turn the working V1 scoring prototype into a simple player, table, and match-record product without overbuilding.",
        size=12,
        color=MUTED,
        after=12,
    )
    for label, value in [
        ("Prepared for", "Cody Butler / Pongle"),
        ("Stack", "Swift iPhone/iPad + Apple Watch app, Supabase backend"),
        ("Date", "May 18, 2026"),
        ("Recommendation", "Build V2 in phases around one question: what lets Cody play ping-pong against somebody new tomorrow?"),
    ]:
        p = add_para(doc, after=2)
        r = p.add_run(f"{label}: ")
        set_run_font(r, bold=True)
        r = p.add_run(value)
        set_run_font(r)


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    add_para(doc, "", after=2)


def fill_table(table, headers, rows, widths):
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, BLUE_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.5, bold=True, color=DARK_BLUE)

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(text)
            set_run_font(r, size=9.2, color=INK)
            if col_idx in (0, 3):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_table_geometry(table, widths)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.text = "Pongle V2 Scope"
    header.paragraph_format.space_after = Pt(0)
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = "Draft scope for discussion - not a fixed-price quote"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED

    add_title_block(doc)

    add_heading(doc, "1. Product Direction", level=1)
    add_callout(
        doc,
        "Recommended V2: keep the live scoring experience from V1, then add the minimum online layer needed for users, tables, match history, and light trust around results.",
    )
    add_para(
        doc,
        "The client is not asking for a full social network yet. The strongest V2 is a practical bridge from \"I kept score in a game\" to \"I know who played, where they played, what the result was, and how that affects player history.\" Supabase should be used for accounts, tables, matches, games, result confirmation, and simple rating data. Swift remains the app layer for iPhone, iPad, and Apple Watch.",
    )

    add_heading(doc, "2. V2 Feature Ranking", level=1)
    headers = ["Rank", "Feature", "Practical V2 scope", "Difficulty / complexity"]
    rows = [
        (
            "1",
            "User profiles",
            "One profile per person with name, phone/email, photo optional, and basic contact invite flow. Phone can help identify people, but consent and duplicate handling must be designed carefully.",
            "Med - auth/data",
        ),
        (
            "2",
            "Table locations",
            "Let users add known tables with venue name, address/map location, access notes, number of tables, and public/private visibility. This directly solves the \"where can I play tomorrow?\" problem.",
            "Med - maps/data",
        ),
        (
            "3",
            "Match and game records",
            "Save completed matches from the V1 scoring flow: players, table, date/time, game scores, match format, server/side info, and input source notes when useful.",
            "Med - data sync",
        ),
        (
            "4",
            "Opponent invite and unclaimed profiles",
            "A user can enter an opponent who is not on Pongle yet. The opponent receives an invite and can claim the history later. This is important, but should be simple in V2.",
            "High - consent/SMS",
        ),
        (
            "5",
            "Result confirmation",
            "Borrow the useful idea from RacketPal: after a match, both players can confirm the result. If the opponent does not respond, mark it as pending or auto-confirm after a clear rule.",
            "High - trust logic",
        ),
        (
            "6",
            "Basic Pongle rating",
            "Start with a simple provisional rating based only on confirmed results. Avoid promising a perfect official rating system in V2. Show rating as early/experimental.",
            "High - rating model",
        ),
        (
            "7",
            "Shareable match cards",
            "Generate an image/card after a match with players, table/location, date, score, and result. Use the iOS share sheet first; do not build automatic social posting yet.",
            "Med - UI/export",
        ),
        (
            "8",
            "Club/table group basics",
            "Allow a simple club or venue label for tables and match history. Keep admin tools light. Full club management can wait.",
            "High - roles/admin",
        ),
        (
            "9",
            "V1 gameplay polish",
            "Keep improving V1 reliability where it supports V2: 3/11/21 point modes, Best of 1/3/5/7 sets, iPad scoreboard clarity, Flic/Watch reliability, and announcement options.",
            "Low-Med - app polish",
        ),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=4)
    fill_table(table, headers, rows, [650, 1750, 5480, 1480])

    add_heading(doc, "3. Recommended Build Order", level=1)
    add_para(
        doc,
        "Phase 1 should prove the data loop: create profile, add table, record a match, and see history. Phase 2 can add trust and growth: opponent confirmation, invite flow, basic rating, and share cards. Phase 3 should wait until testing proves demand: clubs, brackets, payments, subscriptions, and multi-sport support.",
    )

    add_heading(doc, "4. Simple Data Model", level=1)
    model_table = doc.add_table(rows=1 + 6, cols=3)
    model_rows = [
        ("Profiles", "One person/player", "Name, phone/email, photo, home area, claimed/unclaimed status"),
        ("Tables", "Places to play", "Venue, address/map point, access notes, table count, added by"),
        ("Matches", "One match session", "Players, table, date/time, match format, winner, status"),
        ("Games", "Individual games inside a match", "Game number, scores, first server, side order"),
        ("Confirmations", "Trust layer", "Who confirmed, when, disputed/pending/confirmed"),
        ("Clubs", "Later grouping", "SDSU club or venue groups, kept optional for V2"),
    ]
    fill_table(model_table, ["Model", "Plain meaning", "Key details"], model_rows, [1700, 2300, 5360])

    add_heading(doc, "5. Keep Out of V2 Unless Budget Allows", level=1)
    add_para(
        doc,
        "Do not make V2 carry every future idea. Tournament brackets, full club administration, official-grade anti-fraud, paid credits/subscriptions, AirPods input, dedicated TV apps, and pickleball/multi-sport support should be treated as later phases. They are good ideas, but adding them now would make the first online version slower, riskier, and more expensive.",
    )

    add_heading(doc, "6. Acceptance Criteria", level=1)
    add_para(
        doc,
        "A realistic V2 is successful if Cody can add a real table, invite or select a real opponent, play a match using the existing scorekeeper flow, save the result, see match history, and optionally share a clean result card. That is enough to continue real-world testing without overbuilding.",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
