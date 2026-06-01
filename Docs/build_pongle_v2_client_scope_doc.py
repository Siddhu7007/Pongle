from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Docs/Pongle_V2_Client_Scope.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(34, 34, 34)
MUTED = RGBColor(89, 89, 89)
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def dxa_to_inches(dxa):
    return Inches(dxa / 1440)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = dxa_to_inches(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        size, color, before, after = 16, BLUE, 16, 8
    elif level == 2:
        size, color, before, after = 13, BLUE, 12, 6
    else:
        size, color, before, after = 12, DARK_BLUE, 8, 4
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=True)


def add_labeled_line(doc, label, value):
    p = add_para(doc, after=2)
    r = p.add_run(f"{label}: ")
    set_run_font(r, bold=True)
    r = p.add_run(value)
    set_run_font(r)


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    add_para(doc, "", after=2)


def fill_table(table, headers, rows, widths):
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, BLUE_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.4, bold=True, color=DARK_BLUE)

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(text)
            set_run_font(r, size=9.1, color=INK)
            if col_idx in (0, 3):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(table, widths)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.text = "Pongle V2 Scope Draft"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = "Draft for discussion"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED

    add_para(doc, "Pongle V2 Scope Draft", size=23, bold=True, after=4)
    add_para(
        doc,
        "A realistic next step from the V1 scoring prototype toward player profiles, table locations, match records, and early ratings.",
        size=12,
        color=MUTED,
        after=12,
    )
    add_labeled_line(doc, "Prepared for", "Cody Butler / Pongle")
    add_labeled_line(doc, "Recommended stack", "Swift iPhone/iPad + Apple Watch app with Supabase for accounts, data, and sync")
    add_labeled_line(doc, "Working principle", "Build only what helps a user find a table, play a match, record the result, and use that history to find better future games.")

    add_heading(doc, "1. Recommended V2 Direction")
    add_callout(
        doc,
        "V2 should not try to become a full social network, club platform, tournament system, and payment product all at once. The best next version is a focused online layer on top of the scoring experience that already works.",
    )
    add_para(
        doc,
        "The V1 prototype proved that the core scoring loop is interesting: Flic button input, Apple Watch input, iPhone/iPad scoreboard, match flow, and real-game testing. V2 should now answer the next practical question: after a game is played, can Pongle remember who played, where they played, what happened, and use that history to help players find better games later?",
    )

    add_heading(doc, "2. Proposed V2 Feature Scope")
    headers = ["Rank", "Feature", "What it should include", "Difficulty / complexity"]
    rows = [
        ("1", "Player profiles", "One profile per player with name, phone/email, optional photo, and basic history. Keep this simple at first.", "Med - auth/data"),
        ("2", "Table locations", "Users can add known tables, including venue name, address/map point, access notes, and whether the table is public, semi-public, or private.", "Med - maps/data"),
        ("3", "Match and game records", "Save completed matches from the scoring flow: players, table, date/time, game scores, match format, and winner.", "Med - data sync"),
        ("4", "Invite/claim player flow", "If an opponent is not on Pongle yet, create an unclaimed profile and invite them to claim it later. This avoids blocking real play.", "High - consent/SMS"),
        ("5", "Result confirmation", "Both players can confirm a match result. Pending/disputed results should be visible instead of silently trusted.", "High - trust logic"),
        ("6", "Basic Pongle rating", "Start with a simple provisional rating based on confirmed match results. Label it clearly as early/experimental.", "High - rating model"),
        ("7", "Shareable match cards", "After a match, create a clean image/card with players, table, date, and score. Use the iOS share sheet first.", "Med - UI/export"),
        ("8", "Club or venue labels", "Allow tables and matches to be grouped under a club or venue like SDSU Ping-Pong Club. Keep admin tools light.", "High - roles/admin"),
        ("9", "V1 gameplay polish", "Carry forward useful scoring improvements: 3/11/21 points, Best of 1/3/5/7, iPad readability, Flic/Watch reliability, and announcement options.", "Low-Med - polish"),
    ]
    fill_table(doc.add_table(rows=1 + len(rows), cols=4), headers, rows, [650, 1700, 5480, 1530])

    add_heading(doc, "3. Suggested Build Phases")
    add_para(
        doc,
        "Phase 1: accounts/profiles, table locations, and saving match history. This proves the basic data loop and gives real testers something useful immediately.",
    )
    add_para(
        doc,
        "Phase 2: opponent invites, result confirmation, basic rating, and shareable match cards. This adds trust, growth, and the first version of the rating idea.",
    )
    add_para(
        doc,
        "Phase 3: club management, tournament brackets, paid credits/subscriptions, TV display mode, AirPods input, and multi-sport support. These are good ideas, but they should wait until the core V2 loop is proven.",
    )

    add_heading(doc, "4. Simple Data Model")
    model_rows = [
        ("Profiles", "One person/player", "Name, phone/email, photo, claimed/unclaimed status"),
        ("Tables", "Places to play", "Venue, address/map point, access notes, table count, visibility"),
        ("Matches", "One match session", "Players, table, date/time, winner, match format, confirmation status"),
        ("Games", "Games inside a match", "Game number, scores, first server, side order"),
        ("Confirmations", "Trust layer", "Who confirmed, when, pending/disputed/confirmed"),
        ("Clubs/venues", "Optional grouping", "Useful for SDSU or local venue groups, but not the center of V2"),
    ]
    fill_table(doc.add_table(rows=1 + len(model_rows), cols=3), ["Model", "Plain meaning", "Key details"], model_rows, [1700, 2300, 5360])

    add_heading(doc, "5. Success Criteria")
    add_para(
        doc,
        "V2 is successful if a user can create a profile, add or select a real table, play a match using the existing scoring flow, save the result, see match history, and optionally share a match card. That gives Pongle enough structure for real-world testing without overbuilding the whole long-term platform too early.",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
